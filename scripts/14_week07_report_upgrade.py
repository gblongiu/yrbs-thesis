from __future__ import annotations

import argparse
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List

import pandas as pd
from docx import Document

from week06 import PROJECT_ROOT, require_paths


SECTION_ORDER = [
    "Original Plan for This Week",
    "Tasks Accomplished This Week",
    "Comparison of Planned vs Actual",
    "Self-Rating of Progress",
    "Plan for Next Week",
    "Evidence of Tasks Completed",
]

CHANGE_TRACKING_COLUMNS = [
    "feature_name",
    "baseline_importance",
    "baseline_rank",
    "full_importance",
    "full_rank",
    "delta_importance_full_minus_baseline",
    "delta_rank_full_minus_baseline",
    "baseline_source_path",
    "full_source_path",
]


def _fmt(x: float) -> str:
    return f"{x:.6f}"


def _extract_section_lines(md_text: str) -> Dict[str, List[str]]:
    sections: Dict[str, List[str]] = {}
    current = ""
    for line in md_text.splitlines():
        if line.startswith("## "):
            current = line[3:].strip()
            sections[current] = []
            continue
        if current:
            sections[current].append(line)
    return sections


def _validate_md_section_order(md_text: str) -> None:
    headings = re.findall(r"^##\s+(.+)$", md_text, flags=re.MULTILINE)
    observed = headings[: len(SECTION_ORDER)]
    if observed != SECTION_ORDER:
        raise RuntimeError(f"Markdown section order mismatch: {observed}")


def _validate_docx_section_order(docx_path: Path) -> None:
    doc = Document(docx_path)
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
    observed = headings[: len(SECTION_ORDER)]
    if observed != SECTION_ORDER:
        raise RuntimeError(f"DOCX section order mismatch: {observed}")


def _validate_report_constraints(md_path: Path, docx_path: Path) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    if ";" in md_text:
        raise RuntimeError("Markdown report contains semicolon")
    if "—" in md_text:
        raise RuntimeError("Markdown report contains em dash")

    doc = Document(docx_path)
    text_chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_chunks.append(cell.text)
    doc_text = "\n".join(text_chunks)

    if ";" in doc_text:
        raise RuntimeError("DOCX report contains semicolon")
    if "—" in doc_text:
        raise RuntimeError("DOCX report contains em dash")


def _build_feature_importance_change_tracking(project_root: Path) -> Path:
    baseline_path = project_root / "outputs/tables/hgb_seed2026_baseline_perm_importance_summary.csv"
    full_path = project_root / "outputs/tables/hgb_seed2026_full_perm_importance_summary.csv"
    out_path = project_root / "outputs/tables/week06_feature_importance_change_tracking_seed2026.csv"

    require_paths([baseline_path, full_path])

    baseline_df = pd.read_csv(baseline_path)
    full_df = pd.read_csv(full_path)

    baseline_required = {"feature_name", "importance_mean", "mean_rank"}
    full_required = {"feature_name", "mean_importance_neg_brier_score", "mean_rank_neg_brier_score"}

    missing_baseline = sorted(baseline_required - set(baseline_df.columns))
    if missing_baseline:
        raise RuntimeError(f"Baseline importance summary missing columns: {missing_baseline}")

    missing_full = sorted(full_required - set(full_df.columns))
    if missing_full:
        raise RuntimeError(f"Full importance summary missing columns: {missing_full}")

    baseline_small = baseline_df[["feature_name", "importance_mean", "mean_rank"]].rename(
        columns={
            "importance_mean": "baseline_importance",
            "mean_rank": "baseline_rank",
        }
    )
    full_small = full_df[["feature_name", "mean_importance_neg_brier_score", "mean_rank_neg_brier_score"]].rename(
        columns={
            "mean_importance_neg_brier_score": "full_importance",
            "mean_rank_neg_brier_score": "full_rank",
        }
    )

    merged = full_small.merge(baseline_small, on="feature_name", how="left")
    merged["delta_importance_full_minus_baseline"] = merged["full_importance"] - merged["baseline_importance"]
    merged["delta_rank_full_minus_baseline"] = merged["full_rank"] - merged["baseline_rank"]
    merged["baseline_source_path"] = "outputs/tables/hgb_seed2026_baseline_perm_importance_summary.csv"
    merged["full_source_path"] = "outputs/tables/hgb_seed2026_full_perm_importance_summary.csv"

    merged = merged[CHANGE_TRACKING_COLUMNS].sort_values(["full_rank", "feature_name"], kind="mergesort").reset_index(drop=True)

    if merged.empty:
        raise RuntimeError("Feature-importance change tracking table has zero rows")

    missing_bullying = sorted({"x_qn24", "x_qn25"} - set(merged["feature_name"].astype(str).tolist()))
    if missing_bullying:
        raise RuntimeError(f"Required bullying features missing from change tracking table: {missing_bullying}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    merged.to_csv(out_path, index=False)
    return out_path


def _build_markdown(
    *,
    week6_full: pd.DataFrame,
    week6_ablation: pd.DataFrame,
    week6_calibration: pd.DataFrame,
    change_tracking_rel: str,
    equals_baseline: bool,
) -> str:
    held_hgb = week6_full[(week6_full["split_scope"] == "heldout_test") & (week6_full["model"] == "hgb")].iloc[0]
    held_ab = week6_ablation[week6_ablation["split_scope"] == "heldout_test"].copy()

    def _ab_delta(metric: str) -> float:
        return float(held_ab[held_ab["metric"] == metric].iloc[0]["delta_b_minus_a"])

    cal_held = week6_calibration[week6_calibration["split_scope"] == "heldout_test"].copy()
    platt_held = cal_held[cal_held["calibration_method"] == "platt"].iloc[0]
    isotonic_held = cal_held[cal_held["calibration_method"] == "isotonic"].iloc[0]

    evidence_paths = [
        "docs/status_reports/report_03/contract_extracts.md",
        "docs/status_reports/report_03/protocol_lock_confirmation.md",
        "docs/status_reports/report_03/feature_set_definitions.md",
        "outputs/tables/week06_full_feature_comparison_seed2026.csv",
        "outputs/tables/week06_bullying_ablation_comparison_seed2026.csv",
        "outputs/tables/week06_calibration_sensitivity_seed2026.csv",
        "outputs/tables/hgb_seed2026_full_perm_importance_by_fold.csv",
        "outputs/tables/hgb_seed2026_full_perm_importance_summary.csv",
        change_tracking_rel,
        "outputs/metrics/metrics_test_seed2026_hgb_full_none.csv",
        "outputs/metrics/metrics_test_seed2026_hgb_full_platt.csv",
        "outputs/metrics/metrics_test_seed2026_hgb_full_isotonic.csv",
        "outputs/figures/week06_calibration_sensitivity_seed2026.png",
        "scripts/07_week06_pipeline.py",
        "scripts/08_week06_report_package.py",
        "scripts/09_multiseed_stability.py",
        "scripts/10_bootstrap_ci.py",
    ]

    lines: List[str] = []
    lines.append("# Project Status Report 03 Submission")
    lines.append("")
    lines.append("Week Number: Week 6")
    lines.append(f"Date: {datetime.now(timezone.utc).date().isoformat()}")
    lines.append("Student Name: Gabriel Long")
    lines.append("")

    lines.append("## Original Plan for This Week")
    lines.append("- Run full-feature comparison under the frozen Week 4 protocol using seed 2026. Evidence: `outputs/splits/holdout_seed2026.npz`, `outputs/tables/week06_full_feature_comparison_seed2026.csv`.")
    lines.append("- Run bullying-block ablation to quantify incremental predictive value of QN24 and QN25. Evidence: `outputs/tables/week06_bullying_ablation_comparison_seed2026.csv`.")
    lines.append("- Evaluate deferred calibration methods under leakage-safe training-only calibration. Evidence: `outputs/tables/week06_calibration_sensitivity_seed2026.csv`.")
    lines.append("- Maintain reproducibility and auditability under frozen artifact protections. Evidence: `docs/status_reports/report_03/protocol_lock_confirmation.md`.")
    lines.append("")

    lines.append("## Tasks Accomplished This Week")
    lines.append("- Completed full-feature and ablation comparisons with held-out and cross-validation reporting under seed 2026. Evidence: `outputs/tables/week06_full_feature_comparison_seed2026.csv`, `outputs/tables/week06_bullying_ablation_comparison_seed2026.csv`.")
    lines.append("- Completed calibration sensitivity checks for none, Platt, and isotonic using leakage-safe training-only calibration fitting. Evidence: `outputs/tables/week06_calibration_sensitivity_seed2026.csv`.")
    lines.append("- Completed fold-level and summary permutation-importance stability diagnostics for the full HGB model. Evidence: `outputs/tables/hgb_seed2026_full_perm_importance_by_fold.csv`, `outputs/tables/hgb_seed2026_full_perm_importance_summary.csv`.")
    lines.append(f"- Added an explicit feature-importance change tracking artifact across baseline and full feature sets without rerunning models. Evidence: `{change_tracking_rel}`.")
    lines.append("")

    lines.append("## Comparison of Planned vs Actual")
    lines.append("| Planned Week 6 Task | Actual Completion and Evidence |")
    lines.append("| --- | --- |")
    lines.append("| Full-feature model comparison | Completed with held-out and CV outputs under the frozen split protocol. Evidence: `outputs/metrics/metrics_test_seed2026_hgb_full_none.csv`, `outputs/tables/week06_full_feature_comparison_seed2026.csv`. |")
    lines.append("| Bullying-block ablation | Completed with explicit held-out deltas for ranking and calibration metrics. Evidence: `outputs/tables/week06_bullying_ablation_comparison_seed2026.csv`. |")
    lines.append("| Calibration sensitivity | Completed for none, Platt, and isotonic with side-by-side held-out shift reporting. Evidence: `outputs/tables/week06_calibration_sensitivity_seed2026.csv`. |")
    lines.append("")
    lines.append(f"- Held-out full-feature ROC AUC delta versus frozen baseline is `{_fmt(float(held_hgb['delta_roc_auc_vs_logreg_baseline']))}` and held-out Brier delta is `{_fmt(float(held_hgb['delta_brier_vs_logreg_baseline']))}` from `outputs/tables/week06_full_feature_comparison_seed2026.csv`.")
    lines.append(f"- Held-out ablation deltas for HGB full minus bullying to HGB full are ROC AUC `{_fmt(_ab_delta('roc_auc'))}`, PR AUC `{_fmt(_ab_delta('pr_auc'))}`, and Brier `{_fmt(_ab_delta('brier'))}` from `outputs/tables/week06_bullying_ablation_comparison_seed2026.csv`.")
    lines.append(
        f"- Side-by-side held-out calibration comparison in `outputs/tables/week06_calibration_sensitivity_seed2026.csv` reports Platt `delta_brier_vs_none={_fmt(float(platt_held['delta_brier_vs_none']))}` and `delta_calibration_intercept_vs_none={_fmt(float(platt_held['delta_calibration_intercept_vs_none']))}`, and isotonic `delta_brier_vs_none={_fmt(float(isotonic_held['delta_brier_vs_none']))}` and `delta_calibration_intercept_vs_none={_fmt(float(isotonic_held['delta_calibration_intercept_vs_none']))}`."
    )
    lines.append(
        f"- Stability diagnostics inform model selection by showing consistent fold-level importance patterns and transparent cross-feature-set shifts. This reduces dependence on a single split and supports choosing models that are both accurate and stable for interpretation. Evidence: `outputs/tables/hgb_seed2026_full_perm_importance_summary.csv`, `outputs/tables/hgb_seed2026_full_perm_importance_by_fold.csv`, `{change_tracking_rel}`."
    )
    if equals_baseline:
        lines.append("- Structural limitation note: under current modeling-table scope, full-minus-bullying covariates equal the baseline covariates. Evidence: `docs/status_reports/report_03/feature_set_definitions.md`.")
    lines.append("- Scope control statement: Week 6 outputs remained within the frozen protocol and did not modify Week 4 or Week 5 artifacts. Evidence: `docs/status_reports/report_03/protocol_lock_confirmation.md`.")
    lines.append("")

    lines.append("## Self-Rating of Progress")
    lines.append("- [x] Met the Planned Tasks")
    lines.append("- [ ] Did Not Meet the Planned Tasks")
    lines.append("- Rationale: Week 6 planned deliverables were completed and documented with auditable evidence paths and frozen-protocol continuity checks. Evidence: `docs/status_reports/report_03/protocol_lock_confirmation.md`, `docs/status_reports/report_03/rubric_audit.md`.")
    lines.append("")

    lines.append("## Plan for Next Week")
    lines.append("- Execute Week 7 robustness work package using the prepared scripts and frozen Week 6 evidence baseline. Evidence anchor: `scripts/09_multiseed_stability.py`, `scripts/10_bootstrap_ci.py`, `scripts/11_hyperparameter_sensitivity.py`, `scripts/12_subgroup_audit.py`.")
    lines.append("- Extend subgroup and governance analyses as planned next-week deliverables without altering frozen Week 4 and Week 5 artifacts. Evidence anchor: `scripts/13_upgrade_integrity_check.py`, `scripts/14_week07_report_upgrade.py`.")
    lines.append("- Keep non-causal reporting boundaries and calibration-focused interpretation in all next-week outputs. Evidence anchor: `outputs/tables/week06_calibration_sensitivity_seed2026.csv`.")
    lines.append("")

    lines.append("## Evidence of Tasks Completed")
    lines.append("- Repository-relative artifact paths used for this Week 6 submission:")
    for path in evidence_paths:
        lines.append(f"- `{path}`")

    return "\n".join(lines)


def _write_docx(md_text: str, out_path: Path) -> None:
    doc = Document()

    doc.add_paragraph("INFO-I 492 Senior Thesis")
    doc.add_paragraph("Project Status Report #03")
    doc.add_paragraph("Week 6")
    doc.add_paragraph("")

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Light Shading Accent 1"
    meta.cell(0, 0).text = "Student"
    meta.cell(0, 1).text = "Gabriel Long"
    meta.cell(1, 0).text = "Course"
    meta.cell(1, 1).text = "INFO-I 492 Senior Thesis"
    meta.cell(2, 0).text = "Instructor"
    meta.cell(2, 1).text = "Dr. Sridhar Ramachandran"
    meta.cell(3, 0).text = "Report Date"
    meta.cell(3, 1).text = datetime.now(timezone.utc).date().isoformat()

    sections = _extract_section_lines(md_text)

    for section in SECTION_ORDER:
        doc.add_heading(section, level=1)
        lines = sections.get(section, [])

        if section == "Evidence of Tasks Completed":
            section_text = "\n".join(lines)
            paths = re.findall(r"`((?:outputs|docs|scripts)/[^`]+)`", section_text)
            deduped_paths: List[str] = []
            seen = set()
            for p in paths:
                if p not in seen:
                    seen.add(p)
                    deduped_paths.append(p)
            doc.add_paragraph("Evidence paths for submission traceability:")
            for p in deduped_paths:
                doc.add_paragraph(p, style="List Bullet")
            continue

        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("|"):
                doc.add_paragraph(line)
            else:
                doc.add_paragraph(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate Week 6 submission-scoped Status Report 03 artifacts")
    parser.add_argument("--project-root", type=Path, default=PROJECT_ROOT)
    parser.add_argument("--seed", type=int, default=2026)
    args = parser.parse_args()

    if args.seed != 2026:
        raise RuntimeError("Frozen protocol seed must remain 2026")

    report_dir = args.project_root / "docs/status_reports/report_03"
    md_path = report_dir / "Project_Status_Report_03_Submission.md"
    docx_path = report_dir / "Project_Status_Report_03_Submission.docx"

    full_table = args.project_root / "outputs/tables/week06_full_feature_comparison_seed2026.csv"
    ablation_table = args.project_root / "outputs/tables/week06_bullying_ablation_comparison_seed2026.csv"
    calibration_table = args.project_root / "outputs/tables/week06_calibration_sensitivity_seed2026.csv"
    perm_by_fold = args.project_root / "outputs/tables/hgb_seed2026_full_perm_importance_by_fold.csv"
    perm_summary = args.project_root / "outputs/tables/hgb_seed2026_full_perm_importance_summary.csv"
    context_path = report_dir / "week06_context.json"

    require_paths([md_path, docx_path, full_table, ablation_table, calibration_table, perm_by_fold, perm_summary, context_path])

    change_table_path = _build_feature_importance_change_tracking(args.project_root)
    change_table_rel = str(change_table_path.relative_to(args.project_root))

    week6_full = pd.read_csv(full_table)
    week6_ablation = pd.read_csv(ablation_table)
    week6_calibration = pd.read_csv(calibration_table)

    context = json.loads(context_path.read_text(encoding="utf-8"))
    feature_sets = context.get("feature_sets", {}) if isinstance(context, dict) else {}
    equals_baseline = bool(feature_sets.get("equals_baseline", False))

    md_text = _build_markdown(
        week6_full=week6_full,
        week6_ablation=week6_ablation,
        week6_calibration=week6_calibration,
        change_tracking_rel=change_table_rel,
        equals_baseline=equals_baseline,
    )

    _validate_md_section_order(md_text)
    md_path.write_text(md_text, encoding="utf-8")

    _write_docx(md_text, docx_path)

    _validate_docx_section_order(docx_path)
    _validate_report_constraints(md_path, docx_path)

    print(
        json.dumps(
            {
                "status": "ok",
                "markdown": str(md_path),
                "docx": str(docx_path),
                "change_tracking_table": change_table_rel,
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
