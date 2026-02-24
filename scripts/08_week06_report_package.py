from __future__ import annotations

import argparse
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Tuple

import pandas as pd
from docx import Document

from week06 import PROJECT_ROOT, REQUIRED_METRIC_COLS, sha256_file


SECTION_ORDER = [
    "Original Plan for This Week",
    "Tasks Accomplished This Week",
    "Comparison of Planned vs Actual",
    "Self-Rating of Progress",
    "Plan for Next Week",
    "Evidence of Tasks Completed",
]


def _rel(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(PROJECT_ROOT.resolve()))
    except Exception:
        return str(path)


def _fmt(x: float) -> str:
    return f"{x:.6f}"


def _load_context(report_dir: Path) -> Dict[str, object]:
    context_path = report_dir / "week06_context.json"
    if not context_path.exists():
        raise RuntimeError(f"Missing context file: {context_path}")
    return json.loads(context_path.read_text(encoding="utf-8"))


def _check_required_metrics_files(metrics_files: List[Path]) -> List[str]:
    failures: List[str] = []
    for p in metrics_files:
        if not p.exists():
            failures.append(f"Missing metric file: {p}")
            continue
        df = pd.read_csv(p)
        missing = [c for c in REQUIRED_METRIC_COLS if c not in df.columns]
        if missing:
            failures.append(f"Metric file missing required columns {missing}: {p}")
    return failures


def _extract_section_bodies(markdown_text: str) -> Dict[str, str]:
    section_bodies: Dict[str, str] = {}
    headings = list(re.finditer(r"^##\s+(.+)$", markdown_text, flags=re.MULTILINE))
    for i, h in enumerate(headings):
        name = h.group(1).strip()
        start = h.end()
        end = headings[i + 1].start() if i + 1 < len(headings) else len(markdown_text)
        section_bodies[name] = markdown_text[start:end].strip()
    return section_bodies


def _validate_markdown_section_order(markdown_text: str) -> Tuple[bool, List[str]]:
    headings = re.findall(r"^##\s+(.+)$", markdown_text, flags=re.MULTILINE)
    first_six = headings[: len(SECTION_ORDER)]
    if first_six == SECTION_ORDER:
        return True, []
    return False, [f"Markdown section order mismatch. Found={first_six} Expected={SECTION_ORDER}"]


def _validate_docx_section_order(docx_path: Path) -> Tuple[bool, List[str]]:
    doc = Document(docx_path)
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
    first_six = headings[: len(SECTION_ORDER)]
    if first_six == SECTION_ORDER:
        return True, []
    return False, [f"DOCX section order mismatch. Found={first_six} Expected={SECTION_ORDER}"]


def _validate_report_writing_constraints(md_path: Path, docx_path: Path) -> List[str]:
    failures: List[str] = []

    md_text = md_path.read_text(encoding="utf-8")
    if ";" in md_text:
        failures.append("Markdown report contains semicolon characters.")
    if "—" in md_text:
        failures.append("Markdown report contains em dash characters.")

    doc = Document(docx_path)
    doc_text_chunks: List[str] = []
    doc_text_chunks.extend([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                doc_text_chunks.append(cell.text)
    doc_text = "\n".join(doc_text_chunks)

    if ";" in doc_text:
        failures.append("DOCX report contains semicolon characters.")
    if "—" in doc_text:
        failures.append("DOCX report contains em dash characters.")

    return failures


def _ensure_evidence_paths_present(markdown_text: str) -> List[str]:
    failures: List[str] = []
    bodies = _extract_section_bodies(markdown_text)
    for section in SECTION_ORDER:
        body = bodies.get(section, "")
        if not re.search(r"(?:outputs|docs|scripts)/", body):
            failures.append(f"Section '{section}' is missing repository-relative evidence paths.")
    return failures


def build_report_markdown(
    *,
    seed: int,
    feature_sets: Dict[str, object],
    full_df: pd.DataFrame,
    ablation_df: pd.DataFrame,
    calibration_df: pd.DataFrame,
    context: Dict[str, object],
) -> str:
    held_hgb = full_df[(full_df["split_scope"] == "heldout_test") & (full_df["model"] == "hgb")].iloc[0]
    cv_hgb = full_df[(full_df["split_scope"] == "cv_mean") & (full_df["model"] == "hgb")].iloc[0]

    held_ab = ablation_df[ablation_df["split_scope"] == "heldout_test"].copy()
    cv_ab = ablation_df[ablation_df["split_scope"] == "cv_mean"].copy()

    held_cal = calibration_df[calibration_df["split_scope"] == "heldout_test"].copy()
    cv_cal = calibration_df[calibration_df["split_scope"] == "cv_mean"].copy()

    delta_held_roc = float(held_hgb["delta_roc_auc_vs_logreg_baseline"])
    delta_held_brier = float(held_hgb["delta_brier_vs_logreg_baseline"])

    def ab_delta(df: pd.DataFrame, metric: str) -> float:
        return float(df[df["metric"] == metric].iloc[0]["delta_b_minus_a"])

    held_ab_roc = ab_delta(held_ab, "roc_auc")
    held_ab_pr = ab_delta(held_ab, "pr_auc")
    held_ab_brier = ab_delta(held_ab, "brier")

    platt_held = held_cal[held_cal["calibration_method"] == "platt"].iloc[0]
    iso_held = held_cal[held_cal["calibration_method"] == "isotonic"].iloc[0]

    equals_baseline = bool(feature_sets.get("equals_baseline", False))

    lines: List[str] = []
    lines.append("# Project Status Report 03 Submission")
    lines.append("")
    lines.append("Week Number: Week 6")
    lines.append(f"Date: {datetime.now(timezone.utc).date().isoformat()}")
    lines.append("Student Name: Gabriel Long")
    lines.append("")

    lines.append("## Original Plan for This Week")
    lines.append("- Run full-feature comparison under the frozen Week 4 protocol using seed 2026. Evidence: `docs/status_reports/report_03/contract_extracts.md`, `outputs/splits/holdout_seed2026.npz`.")
    lines.append("- Run bullying-block ablation to quantify incremental predictive value of QN24 and QN25. Evidence: `docs/status_reports/report_03/contract_extracts.md`.")
    lines.append("- Evaluate deferred calibration methods Platt and isotonic under leakage-safe training-only calibration. Evidence: `docs/status_reports/report_03/contract_extracts.md`.")
    lines.append("- Maintain reproducibility and auditability with frozen artifact checks and additive outputs only. Evidence: `docs/status_reports/report_03/protocol_lock_confirmation.md`.")
    lines.append("")

    lines.append("## Tasks Accomplished This Week")
    lines.append("- Built and executed a Week 6 guardrail-first pipeline with frozen artifact hash checks before and after modeling. Evidence: `scripts/07_week06_pipeline.py`, `docs/status_reports/report_03/protocol_lock_confirmation.md`.")
    lines.append("- Generated required Week 6 metrics files for full-feature, ablation, and calibration sensitivity configurations. Evidence: `outputs/metrics/metrics_cv_seed2026_hgb_full_none.csv`, `outputs/metrics/metrics_test_seed2026_hgb_full_isotonic.csv`.")
    lines.append("- Generated Week 6 comparison tables and figures with explicit deltas and source metric paths. Evidence: `outputs/tables/week06_full_feature_comparison_seed2026.csv`, `outputs/figures/week06_calibration_sensitivity_seed2026.png`.")
    lines.append("- Generated permutation-importance stability outputs with required bullying-feature rows. Evidence: `outputs/tables/hgb_seed2026_full_perm_importance_summary.csv`.")
    lines.append("- Updated run logs and decisions log with commands, feature definitions, calibration protocol, and equality-case handling. Evidence: `docs/experiment_log.md`, `docs/decisions_log.md`.")
    lines.append("")

    lines.append("## Comparison of Planned vs Actual")
    lines.append("| Planned Week 6 Task | Actual Completion and Evidence |")
    lines.append("| --- | --- |")
    lines.append("| Full-feature model comparison under frozen protocol | Completed with CV and held-out outputs for logreg and HGB full feature sets. Evidence: `outputs/metrics/metrics_cv_seed2026_logreg_full_none.csv`, `outputs/metrics/metrics_test_seed2026_hgb_full_none.csv`. |")
    lines.append("| Bullying-block ablation | Completed with explicit delta table. Evidence: `outputs/tables/week06_bullying_ablation_comparison_seed2026.csv`. |")
    lines.append("| Calibration sensitivity with Platt and isotonic | Completed with CV and held-out metrics and sensitivity table. Evidence: `outputs/metrics/metrics_test_seed2026_hgb_full_platt.csv`, `outputs/metrics/metrics_test_seed2026_hgb_full_isotonic.csv`, `outputs/tables/week06_calibration_sensitivity_seed2026.csv`. |")
    lines.append("| Interpretability and stability outputs | Completed with fold-level and summary permutation importance outputs. Evidence: `outputs/tables/hgb_seed2026_full_perm_importance_by_fold.csv`, `outputs/tables/hgb_seed2026_full_perm_importance_summary.csv`. |")
    lines.append("")
    lines.append("Analytical interpretation:")
    lines.append(f"- Full-feature held-out ROC AUC difference versus frozen logreg baseline is `{_fmt(delta_held_roc)}` and held-out Brier difference is `{_fmt(delta_held_brier)}` from `outputs/tables/week06_full_feature_comparison_seed2026.csv`. These results indicate how discrimination and probability quality moved under the research question contrast.")
    lines.append(f"- Bullying ablation held-out deltas are ROC AUC `{_fmt(held_ab_roc)}`, PR AUC `{_fmt(held_ab_pr)}`, and Brier `{_fmt(held_ab_brier)}` from `outputs/tables/week06_bullying_ablation_comparison_seed2026.csv`. This quantifies incremental predictive value out of sample without causal interpretation.")
    lines.append(
        f"- Calibration sensitivity on held-out data shows Platt delta Brier `{_fmt(float(platt_held['delta_brier_vs_none']))}` and isotonic delta Brier `{_fmt(float(iso_held['delta_brier_vs_none']))}` versus none from `outputs/tables/week06_calibration_sensitivity_seed2026.csv`."
    )
    if equals_baseline:
        lines.append("- Ablation control-set note: under current modeling-table scope, `full_minus_bullying_features` equals baseline covariates. This is documented as a structural limitation in `docs/status_reports/report_03/feature_set_definitions.md` and `docs/decisions_log.md`.")
    lines.append("- Scope control statement: Week 6 remained within the proposal timeline and did not alter frozen Week 4 or Week 5 artifacts. Evidence: `docs/status_reports/report_03/protocol_lock_confirmation.md`.")
    lines.append("")

    lines.append("## Self-Rating of Progress")
    lines.append("- [x] Met the Planned Tasks")
    lines.append("- [ ] Did Not Meet the Planned Tasks")
    lines.append("- Rationale: Planned Week 6 deliverables were produced with frozen-protocol continuity and artifact-level audit checks. Evidence: `docs/status_reports/report_03/rubric_audit.md`, `docs/status_reports/report_03/protocol_lock_confirmation.md`.")
    lines.append("")

    lines.append("## Plan for Next Week")
    lines.append("- Finalize Week 7 model selection package using Week 6 comparative evidence. Evidence base: `outputs/tables/week06_full_feature_comparison_seed2026.csv`.")
    lines.append("- Expand presentation-ready calibration interpretation and probability-quality narrative. Evidence base: `outputs/tables/week06_calibration_sensitivity_seed2026.csv`, `outputs/figures/week06_calibration_sensitivity_seed2026.png`.")
    if equals_baseline:
        lines.append("- Expand the modeling table with additional allowable covariates so the ablation control set differs from baseline while preserving leakage safeguards. Evidence anchor: `docs/status_reports/report_03/feature_set_definitions.md`.")
    lines.append("- Preserve non-causal reporting and frozen-protocol reproducibility controls in Week 7 artifacts. Evidence anchor: `docs/status_reports/report_03/contract_extracts.md`.")
    lines.append("")

    lines.append("## Evidence of Tasks Completed")
    lines.append("- Contract and protocol documents: `docs/status_reports/report_03/contract_extracts.md`, `docs/status_reports/report_03/protocol_lock_confirmation.md`, `docs/status_reports/report_03/feature_set_definitions.md`.")
    lines.append("- Week 6 metrics: `outputs/metrics/metrics_cv_seed2026_logreg_full_none.csv`, `outputs/metrics/metrics_test_seed2026_logreg_full_none.csv`, `outputs/metrics/metrics_cv_seed2026_hgb_full_none.csv`, `outputs/metrics/metrics_test_seed2026_hgb_full_none.csv`, `outputs/metrics/metrics_cv_seed2026_hgb_full_minus_bullying.csv`, `outputs/metrics/metrics_test_seed2026_hgb_full_minus_bullying.csv`, `outputs/metrics/metrics_cv_seed2026_hgb_full_platt.csv`, `outputs/metrics/metrics_cv_seed2026_hgb_full_isotonic.csv`, `outputs/metrics/metrics_test_seed2026_hgb_full_platt.csv`, `outputs/metrics/metrics_test_seed2026_hgb_full_isotonic.csv`.")
    lines.append("- Week 6 tables and figures: `outputs/tables/week06_full_feature_comparison_seed2026.csv`, `outputs/tables/week06_bullying_ablation_comparison_seed2026.csv`, `outputs/tables/week06_calibration_sensitivity_seed2026.csv`, `outputs/figures/week06_full_feature_comparison_seed2026.png`, `outputs/figures/week06_bullying_ablation_comparison_seed2026.png`, `outputs/figures/week06_calibration_sensitivity_seed2026.png`.")
    lines.append("- Stability outputs: `outputs/tables/hgb_seed2026_full_perm_importance_by_fold.csv`, `outputs/tables/hgb_seed2026_full_perm_importance_summary.csv`.")
    lines.append("- Logging and audit trail: `docs/experiment_log.md`, `docs/decisions_log.md`, `docs/status_reports/report_03/rubric_audit.md`, `docs/status_reports/report_03/week06_run_manifest.json`.")

    text = "\n".join(lines)
    return text


def build_report_docx(md_text: str, out_docx: Path) -> None:
    doc = Document()
    doc.add_heading("Project Status Report 03 Submission", level=0)

    # Split markdown into sections by heading for deterministic docx structure.
    bodies = _extract_section_bodies(md_text)

    for section in SECTION_ORDER:
        doc.add_heading(section, level=1)
        body = bodies.get(section, "")
        for raw_line in body.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("|"):
                # Keep markdown tables as plain text blocks to avoid conversion complexity.
                doc.add_paragraph(line)
            else:
                doc.add_paragraph(line)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)


def build_manifest(
    *,
    report_dir: Path,
    context: Dict[str, object],
    extra_files: List[Path],
) -> Dict[str, object]:
    feature_sets = context["feature_sets"]

    candidate_paths: List[Path] = [
        PROJECT_ROOT / "scripts" / "week06.py",
        PROJECT_ROOT / "scripts" / "07_week06_pipeline.py",
        PROJECT_ROOT / "scripts" / "08_week06_report_package.py",
        report_dir / "contract_extracts.md",
        report_dir / "protocol_lock_confirmation.md",
        report_dir / "feature_set_definitions.md",
        report_dir / "Project_Status_Report_03_Submission.md",
        report_dir / "Project_Status_Report_03_Submission.docx",
        report_dir / "rubric_audit.md",
        report_dir / "week06_context.json",
    ]

    for rel_group in context["paths"].values():
        if isinstance(rel_group, list):
            for rel in rel_group:
                candidate_paths.append(PROJECT_ROOT / rel)
        elif isinstance(rel_group, str):
            candidate_paths.append(PROJECT_ROOT / rel_group)

    candidate_paths.extend(extra_files)

    # Deduplicate while preserving order.
    seen = set()
    deduped: List[Path] = []
    for p in candidate_paths:
        rp = str(p.resolve())
        if rp in seen:
            continue
        seen.add(rp)
        deduped.append(p)

    missing = [str(p) for p in deduped if not p.exists()]
    if missing:
        raise RuntimeError("Manifest build failed. Missing files:\n" + "\n".join(missing))

    entries = []
    for p in deduped:
        entries.append({"path": _rel(p), "sha256": sha256_file(p)})

    manifest = {
        "timestamp_utc": datetime.now(timezone.utc).isoformat(),
        "seed": int(context["seed"]),
        "baseline_features": feature_sets["baseline_features"],
        "full_features": feature_sets["full_features"],
        "full_minus_bullying_features": feature_sets["full_minus_bullying_features"],
        "new_files": entries,
    }
    return manifest


def main() -> None:
    parser = argparse.ArgumentParser(description="Week 6 report packaging and rubric audit.")
    parser.add_argument("--seed", type=int, default=2026)
    parser.add_argument("--project-root", type=Path, default=PROJECT_ROOT)
    args = parser.parse_args()

    if args.seed != 2026:
        raise RuntimeError("Week 6 report package requires seed 2026.")

    report_dir = args.project_root / "docs" / "status_reports" / "report_03"
    report_dir.mkdir(parents=True, exist_ok=True)

    context = _load_context(report_dir)

    full_table = args.project_root / "outputs" / "tables" / "week06_full_feature_comparison_seed2026.csv"
    ablation_table = args.project_root / "outputs" / "tables" / "week06_bullying_ablation_comparison_seed2026.csv"
    calibration_table = args.project_root / "outputs" / "tables" / "week06_calibration_sensitivity_seed2026.csv"

    full_df = pd.read_csv(full_table)
    ablation_df = pd.read_csv(ablation_table)
    calibration_df = pd.read_csv(calibration_table)

    md_text = build_report_markdown(
        seed=args.seed,
        feature_sets=context["feature_sets"],
        full_df=full_df,
        ablation_df=ablation_df,
        calibration_df=calibration_df,
        context=context,
    )

    md_path = report_dir / "Project_Status_Report_03_Submission.md"
    docx_path = report_dir / "Project_Status_Report_03_Submission.docx"
    md_path.write_text(md_text, encoding="utf-8")
    build_report_docx(md_text, docx_path)

    rubric_checks: List[Tuple[str, bool, str]] = []

    # Section order checks.
    ok_md_order, md_order_failures = _validate_markdown_section_order(md_text)
    rubric_checks.append(("Section order (Markdown)", ok_md_order, " | ".join(md_order_failures) if md_order_failures else "ok"))

    ok_docx_order, docx_order_failures = _validate_docx_section_order(docx_path)
    rubric_checks.append(("Section order (DOCX)", ok_docx_order, " | ".join(docx_order_failures) if docx_order_failures else "ok"))

    metric_files = [args.project_root / rel for rel in context["paths"]["metrics"]]
    metric_failures = _check_required_metrics_files(metric_files)
    rubric_checks.append(("Required metrics present", len(metric_failures) == 0, " | ".join(metric_failures) if metric_failures else "ok"))

    required_artifacts = [
        args.project_root / rel for rel in context["paths"]["tables"] + context["paths"]["figures"]
    ]
    missing_artifacts = [str(p) for p in required_artifacts if not p.exists()]
    rubric_checks.append(("Required artifacts present", len(missing_artifacts) == 0, " | ".join(missing_artifacts) if missing_artifacts else "ok"))

    evidence_failures = _ensure_evidence_paths_present(md_text)
    rubric_checks.append(("Evidence paths present per section", len(evidence_failures) == 0, " | ".join(evidence_failures) if evidence_failures else "ok"))

    writing_failures = _validate_report_writing_constraints(md_path, docx_path)
    rubric_checks.append(("Report writing constraints", len(writing_failures) == 0, " | ".join(writing_failures) if writing_failures else "ok"))

    frozen_ok = bool(context.get("frozen_hash_check_passed", False))
    rubric_checks.append(("Frozen artifact hashes unchanged", frozen_ok, "ok" if frozen_ok else "frozen_hash_check_passed is false"))

    all_ok = all(flag for _, flag, _ in rubric_checks)

    audit_lines = []
    audit_lines.append("# Week 6 Rubric Audit")
    audit_lines.append("")
    for name, flag, detail in rubric_checks:
        audit_lines.append(f"- {'PASS' if flag else 'FAIL'}: {name}")
        audit_lines.append(f"  - Detail: {detail}")
    audit_lines.append("")
    audit_lines.append(f"Overall: {'PASS' if all_ok else 'FAIL'}")

    rubric_path = report_dir / "rubric_audit.md"
    rubric_path.write_text("\n".join(audit_lines), encoding="utf-8")

    manifest = build_manifest(
        report_dir=report_dir,
        context=context,
        extra_files=[rubric_path, md_path, docx_path],
    )

    manifest_path = report_dir / "week06_run_manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2, sort_keys=True), encoding="utf-8")

    # Verify hashes after write to satisfy explicit correctness check.
    loaded_manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    for entry in loaded_manifest["new_files"]:
        p = args.project_root / entry["path"]
        actual = sha256_file(p)
        if actual != entry["sha256"]:
            raise RuntimeError(f"Manifest hash mismatch for {entry['path']}")

    if not all_ok:
        raise RuntimeError("Rubric audit failed. See docs/status_reports/report_03/rubric_audit.md")

    print(json.dumps({"status": "ok", "report_markdown": _rel(md_path), "report_docx": _rel(docx_path)}, indent=2))


if __name__ == "__main__":
    main()
