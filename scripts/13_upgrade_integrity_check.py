from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Dict, List, Sequence, Tuple

import pandas as pd
from docx import Document

from week06 import PROJECT_ROOT, require_paths, sha256_file


REPORT_SECTION_ORDER = [
    "Original Plan for This Week",
    "Tasks Accomplished This Week",
    "Comparison of Planned vs Actual",
    "Self-Rating of Progress",
    "Plan for Next Week",
    "Evidence of Tasks Completed",
]


FROZEN_ARTIFACTS = [
    Path("data/processed/yrbs_2023_modeling.parquet"),
    Path("outputs/splits/holdout_seed2026.npz"),
    Path("outputs/splits/cvfolds_seed2026.npz"),
    Path("outputs/tuning/hgb_seed2026_baseline_best_params.json"),
    Path("outputs/metrics/metrics_cv_seed2026_logreg_baseline_none.csv"),
    Path("outputs/metrics/metrics_test_seed2026_logreg_baseline_none.csv"),
    Path("outputs/metrics/metrics_cv_seed2026_hgb_baseline_none.csv"),
    Path("outputs/metrics/metrics_test_seed2026_hgb_baseline_none.csv"),
]


WEEK6_HASH_WHITELIST = {
    "docs/status_reports/report_03/Project_Status_Report_03_Submission.md",
    "docs/status_reports/report_03/Project_Status_Report_03_Submission.docx",
    "docs/status_reports/report_03/rubric_audit.md",
}


def ensure_required_columns(df: pd.DataFrame, required: Sequence[str], label: str) -> None:
    missing = [col for col in required if col not in df.columns]
    if missing:
        raise RuntimeError(f"{label} missing required columns: {missing}")


def check_week06_manifest_immutability(
    *,
    project_root: Path,
    manifest_path: Path,
    whitelist: Sequence[str],
) -> List[str]:
    whitelist_set = set(whitelist)
    payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    entries = payload.get("new_files", [])
    if not isinstance(entries, list):
        return ["week06 manifest new_files is not a list"]

    failures: List[str] = []
    for entry in entries:
        if not isinstance(entry, dict):
            failures.append("week06 manifest entry has invalid type")
            continue

        rel_path = str(entry.get("path", ""))
        expected = str(entry.get("sha256", ""))
        if not rel_path or not expected:
            failures.append("week06 manifest entry missing path or sha256")
            continue

        if rel_path in whitelist_set:
            continue

        abs_path = project_root / rel_path
        if not abs_path.exists():
            failures.append(f"manifest path missing on disk: {rel_path}")
            continue

        actual = sha256_file(abs_path)
        if actual != expected:
            failures.append(f"manifest hash mismatch for {rel_path}")

    return failures


def extract_repo_paths_from_markdown(md_text: str) -> List[str]:
    pattern = re.compile(r"(?:docs|outputs|scripts|data)/[A-Za-z0-9_./-]+")
    raw = pattern.findall(md_text)
    cleaned: List[str] = []
    for path_str in raw:
        cleaned.append(path_str.rstrip(".,)"))
    return sorted(set(cleaned))


def validate_report_section_order(md_text: str) -> List[str]:
    headings = re.findall(r"^##\s+(.+)$", md_text, flags=re.MULTILINE)
    observed = headings[: len(REPORT_SECTION_ORDER)]
    if observed != REPORT_SECTION_ORDER:
        return [f"report section order mismatch: observed={observed}"]
    return []


def validate_docx_section_order(docx_path: Path) -> List[str]:
    doc = Document(docx_path)
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
    idx = 0
    for heading in headings:
        if idx < len(REPORT_SECTION_ORDER) and heading == REPORT_SECTION_ORDER[idx]:
            idx += 1
    if idx != len(REPORT_SECTION_ORDER):
        return [f"docx section order mismatch: observed={headings}"]
    return []


def check_text_constraints(md_path: Path, docx_path: Path) -> List[str]:
    failures: List[str] = []

    md_text = md_path.read_text(encoding="utf-8")
    if ";" in md_text:
        failures.append("markdown report contains semicolon")
    if "—" in md_text:
        failures.append("markdown report contains em dash")

    doc = Document(docx_path)
    doc_chunks: List[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                doc_chunks.append(cell.text)
    doc_text = "\n".join(doc_chunks)

    if ";" in doc_text:
        failures.append("docx report contains semicolon")
    if "—" in doc_text:
        failures.append("docx report contains em dash")

    return failures


def _expected_hash_for_path(path: Path, mapping: Dict[str, str], project_root: Path) -> str | None:
    candidates = [
        str(path),
        str(path.resolve()),
    ]
    try:
        candidates.append(str(path.resolve().relative_to(project_root.resolve())))
    except Exception:
        pass

    for key in candidates:
        if key in mapping:
            return mapping[key]
    return None


def run_integrity_checks(project_root: Path) -> Tuple[bool, List[str]]:
    messages: List[str] = []

    week7_required = [
        project_root / "outputs/tables/multiseed_stability_seed2026_2029.csv",
        project_root / "outputs/preds/preds_test_hgb_full_none_seed2026.csv",
        project_root / "outputs/tables/heldout_bootstrap_ci_seed2026.csv",
        project_root / "outputs/tables/hgb_hyperparameter_sensitivity_seed2026.csv",
        project_root / "outputs/tables/hgb_seed2026_full_perm_importance_summary_extended.csv",
        project_root / "outputs/tables/subgroup_performance_seed2026.csv",
        project_root / "model_selection_framework.md",
        project_root / "deployment_and_use_constraints.md",
        project_root / "traceability_matrix.csv",
        project_root / "qa_checklist.md",
        project_root / "risk_register.md",
        project_root / "docs/status_reports/report_03/Project_Status_Report_03_Submission.md",
        project_root / "docs/status_reports/report_03/Project_Status_Report_03_Submission.docx",
    ]

    try:
        require_paths(week7_required)
        messages.append("PASS all required Week 7 files exist")
    except Exception as exc:
        messages.append(f"FAIL required Week 7 files: {exc}")

    try:
        multi_df = pd.read_csv(project_root / "outputs/tables/multiseed_stability_seed2026_2029.csv")
        ensure_required_columns(
            multi_df,
            [
                "seed",
                "roc_auc_mean",
                "pr_auc_mean",
                "brier_mean",
                "slope_mean",
                "n_train",
                "cv_folds",
                "test_split_seed",
                "roc_auc_std_across_seeds",
                "pr_auc_std_across_seeds",
                "brier_std_across_seeds",
                "slope_std_across_seeds",
            ],
            "multiseed_stability",
        )
        if len(multi_df) != 4:
            raise RuntimeError("multiseed row count must be exactly 4")
        messages.append("PASS multiseed schema and row count")
    except Exception as exc:
        messages.append(f"FAIL multiseed schema or row count: {exc}")

    try:
        boot_df = pd.read_csv(project_root / "outputs/tables/heldout_bootstrap_ci_seed2026.csv")
        ensure_required_columns(
            boot_df,
            [
                "metric",
                "mean",
                "lower_95",
                "upper_95",
                "bootstrap_seed",
                "n_boot",
                "sample_scope",
                "stratified",
            ],
            "heldout_bootstrap_ci",
        )
        required_metrics = {"roc_auc", "brier", "calibration_slope"}
        observed_metrics = set(boot_df["metric"].astype(str).tolist())
        if observed_metrics != required_metrics:
            raise RuntimeError(f"bootstrap metric set mismatch: {observed_metrics}")
        messages.append("PASS bootstrap schema and metric set")
    except Exception as exc:
        messages.append(f"FAIL bootstrap schema or metric set: {exc}")

    try:
        sens_df = pd.read_csv(project_root / "outputs/tables/hgb_hyperparameter_sensitivity_seed2026.csv")
        ensure_required_columns(
            sens_df,
            [
                "config_name",
                "seed",
                "cv_folds",
                "brier_mean",
                "calibration_slope_mean",
                "roc_auc_mean",
                "pr_auc_mean",
                "calibration_intercept_mean",
                "params_json",
            ],
            "hgb_hyperparameter_sensitivity",
        )
        if len(sens_df) != 7:
            raise RuntimeError("hyperparameter sensitivity row count must be exactly 7")
        messages.append("PASS hyperparameter sensitivity schema and row count")
    except Exception as exc:
        messages.append(f"FAIL hyperparameter sensitivity schema or row count: {exc}")

    try:
        subgroup_df = pd.read_csv(project_root / "outputs/tables/subgroup_performance_seed2026.csv")
        ensure_required_columns(
            subgroup_df,
            [
                "subgroup_var",
                "subgroup_value",
                "n",
                "n_pos",
                "n_neg",
                "roc_auc",
                "brier",
                "calibration_slope",
                "roc_auc_defined_flag",
                "slope_defined_flag",
                "low_slope_flag",
                "high_error_flag",
            ],
            "subgroup_performance",
        )
        observed_vars = set(subgroup_df["subgroup_var"].astype(str).tolist())
        if observed_vars != {"raceeth", "q2"}:
            raise RuntimeError(f"subgroup dimensions mismatch: {observed_vars}")
        messages.append("PASS subgroup schema and dimensions")
    except Exception as exc:
        messages.append(f"FAIL subgroup schema or dimensions: {exc}")

    try:
        perm_ext_df = pd.read_csv(project_root / "outputs/tables/hgb_seed2026_full_perm_importance_summary_extended.csv")
        ensure_required_columns(
            perm_ext_df,
            [
                "feature_name",
                "coefficient_of_variation_neg_brier_score",
                "fold_consistency_flag",
            ],
            "perm_importance_summary_extended",
        )
        names = set(perm_ext_df["feature_name"].astype(str).tolist())
        for feature_name in ["x_qn24", "x_qn25"]:
            if feature_name not in names:
                raise RuntimeError(f"missing required feature in extended permutation summary: {feature_name}")
        messages.append("PASS extended permutation summary checks")
    except Exception as exc:
        messages.append(f"FAIL extended permutation summary checks: {exc}")

    report_md = project_root / "docs/status_reports/report_03/Project_Status_Report_03_Submission.md"
    report_docx = project_root / "docs/status_reports/report_03/Project_Status_Report_03_Submission.docx"

    try:
        md_text = report_md.read_text(encoding="utf-8")
        md_paths = extract_repo_paths_from_markdown(md_text)
        missing_paths = [path for path in md_paths if not (project_root / path).exists()]
        if missing_paths:
            raise RuntimeError(f"missing referenced paths: {missing_paths}")
        messages.append("PASS report referenced paths exist")
    except Exception as exc:
        messages.append(f"FAIL report referenced paths: {exc}")

    try:
        order_failures = validate_report_section_order(report_md.read_text(encoding="utf-8"))
        order_failures.extend(validate_docx_section_order(report_docx))
        if order_failures:
            raise RuntimeError(" | ".join(order_failures))
        messages.append("PASS report section order")
    except Exception as exc:
        messages.append(f"FAIL report section order: {exc}")

    try:
        text_failures = check_text_constraints(report_md, report_docx)
        if text_failures:
            raise RuntimeError(" | ".join(text_failures))
        messages.append("PASS report text constraints")
    except Exception as exc:
        messages.append(f"FAIL report text constraints: {exc}")

    context_path = project_root / "docs/status_reports/report_03/week06_context.json"
    try:
        require_paths([context_path])
        context = json.loads(context_path.read_text(encoding="utf-8"))
        hash_mapping = context.get("frozen_hashes_pre", {})
        if not isinstance(hash_mapping, dict):
            raise RuntimeError("week06_context frozen_hashes_pre missing or invalid")

        for rel_path in FROZEN_ARTIFACTS:
            abs_path = project_root / rel_path
            if not abs_path.exists():
                raise RuntimeError(f"frozen artifact missing: {rel_path}")
            expected_hash = _expected_hash_for_path(abs_path, hash_mapping, project_root)
            if expected_hash is None:
                raise RuntimeError(f"expected hash not found in context for {rel_path}")
            actual_hash = sha256_file(abs_path)
            if actual_hash != expected_hash:
                raise RuntimeError(f"frozen artifact hash changed: {rel_path}")

        messages.append("PASS frozen Week 4 and Week 5 hash set")
    except Exception as exc:
        messages.append(f"FAIL frozen hash set: {exc}")

    manifest_path = project_root / "docs/status_reports/report_03/week06_run_manifest.json"
    try:
        require_paths([manifest_path])
        failures = check_week06_manifest_immutability(
            project_root=project_root,
            manifest_path=manifest_path,
            whitelist=sorted(WEEK6_HASH_WHITELIST),
        )
        if failures:
            raise RuntimeError(" | ".join(failures))
        messages.append("PASS week06 manifest immutability")
    except Exception as exc:
        messages.append(f"FAIL week06 manifest immutability: {exc}")

    ok = all(msg.startswith("PASS") for msg in messages)
    return ok, messages


def main() -> None:
    parser = argparse.ArgumentParser(description="Week 7 upgrade integrity checker")
    parser.add_argument("--project-root", type=Path, default=PROJECT_ROOT)
    parser.add_argument("--outdir", type=Path, default=Path("outputs"))
    args = parser.parse_args()

    audit_path = args.project_root / args.outdir / "audits" / "week07_upgrade_integrity_audit.md"
    audit_path.parent.mkdir(parents=True, exist_ok=True)
    # Create placeholder early so report path existence checks can pass deterministically.
    if not audit_path.exists():
        audit_path.write_text("# Week 7 Upgrade Integrity Audit\n\nStatus: running\n", encoding="utf-8")

    ok, messages = run_integrity_checks(args.project_root)

    lines: List[str] = []
    lines.append("# Week 7 Upgrade Integrity Audit")
    lines.append("")
    for msg in messages:
        lines.append(f"- {msg}")
    lines.append("")
    lines.append(f"Overall: {'PASS' if ok else 'FAIL'}")

    audit_path.write_text("\n".join(lines), encoding="utf-8")
    print(json.dumps({"status": "ok" if ok else "fail", "audit": str(audit_path)}, indent=2))

    if not ok:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
