from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Dict, List, Tuple

import pandas as pd
from docx import Document

from week06 import PROJECT_ROOT, require_paths


SECTION_ORDER = [
    "Original Plan for This Week",
    "Tasks Accomplished This Week",
    "Comparison of Planned vs Actual",
    "Self-Rating of Progress",
    "Plan for Next Week",
    "Evidence of Tasks Completed",
]

REQUIRED_CHANGE_COLUMNS = [
    "feature_name",
    "baseline_importance",
    "baseline_rank",
    "full_importance",
    "full_rank",
    "delta_importance_full_minus_baseline",
    "delta_rank_full_minus_baseline",
    "baseline_source_path",
    "full_source_path",
]

PATH_PATTERN = re.compile(r"(?:outputs|docs|scripts)/[A-Za-z0-9_./-]+")


def _extract_md_sections(md_text: str) -> Dict[str, str]:
    section_bodies: Dict[str, str] = {}
    headings = list(re.finditer(r"^##\s+(.+)$", md_text, flags=re.MULTILINE))
    for i, match in enumerate(headings):
        name = match.group(1).strip()
        start = match.end()
        end = headings[i + 1].start() if i + 1 < len(headings) else len(md_text)
        section_bodies[name] = md_text[start:end].strip()
    return section_bodies


def _validate_md_section_order(md_text: str) -> Tuple[bool, str]:
    headings = re.findall(r"^##\s+(.+)$", md_text, flags=re.MULTILINE)
    observed = headings[: len(SECTION_ORDER)]
    if observed != SECTION_ORDER:
        return False, f"observed={observed} expected={SECTION_ORDER}"
    return True, "ok"


def _validate_docx_section_order(docx_path: Path) -> Tuple[bool, str]:
    doc = Document(docx_path)
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
    observed = headings[: len(SECTION_ORDER)]
    if observed != SECTION_ORDER:
        return False, f"observed={observed} expected={SECTION_ORDER}"
    return True, "ok"


def _validate_text_constraints(md_path: Path, docx_path: Path) -> Tuple[bool, str]:
    md_text = md_path.read_text(encoding="utf-8")
    if ";" in md_text:
        return False, "markdown contains semicolon"
    if "—" in md_text:
        return False, "markdown contains em dash"

    doc = Document(docx_path)
    doc_text_chunks: List[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                doc_text_chunks.append(cell.text)
    doc_text = "\n".join(doc_text_chunks)

    if ";" in doc_text:
        return False, "docx contains semicolon"
    if "—" in doc_text:
        return False, "docx contains em dash"
    return True, "ok"


def _extract_docx_evidence_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    paragraphs = doc.paragraphs

    start = None
    for i, p in enumerate(paragraphs):
        if p.style.name.startswith("Heading 1") and p.text.strip() == "Evidence of Tasks Completed":
            start = i + 1
            break

    if start is None:
        return ""

    collected: List[str] = []
    for j in range(start, len(paragraphs)):
        p = paragraphs[j]
        if p.style.name.startswith("Heading 1"):
            break
        if p.text.strip():
            collected.append(p.text.strip())

    return "\n".join(collected)


def _evidence_path_density_checks(md_path: Path, docx_path: Path) -> Tuple[bool, str]:
    md_text = md_path.read_text(encoding="utf-8")
    sections = _extract_md_sections(md_text)
    md_evidence = sections.get("Evidence of Tasks Completed", "")
    md_paths = sorted(set(PATH_PATTERN.findall(md_evidence)))

    docx_evidence_text = _extract_docx_evidence_text(docx_path)
    docx_paths = sorted(set(PATH_PATTERN.findall(docx_evidence_text)))

    if len(md_paths) < 10:
        return False, f"markdown evidence path count is {len(md_paths)} which is below 10"
    if len(docx_paths) < 10:
        return False, f"docx evidence path count is {len(docx_paths)} which is below 10"
    return True, f"markdown_paths={len(md_paths)} docx_paths={len(docx_paths)}"


def _calibration_table_check(cal_path: Path) -> Tuple[bool, str]:
    if not cal_path.exists():
        return False, f"missing file: {cal_path}"
    df = pd.read_csv(cal_path)
    if "calibration_method" not in df.columns:
        return False, "calibration_method column missing"
    methods = {str(x).strip().lower() for x in df["calibration_method"].tolist()}
    required = {"none", "platt", "isotonic"}
    missing = sorted(required - methods)
    if missing:
        return False, f"missing methods: {missing}"
    return True, "ok"


def _change_tracking_check(path: Path) -> Tuple[bool, str, bool, str]:
    if not path.exists():
        return False, f"missing file: {path}", False, "output table missing"

    df = pd.read_csv(path)
    missing_cols = [c for c in REQUIRED_CHANGE_COLUMNS if c not in df.columns]
    if missing_cols:
        return False, f"missing columns: {missing_cols}", False, "required columns missing"

    feature_names = set(df["feature_name"].astype(str).tolist())
    missing_features = sorted({"x_qn24", "x_qn25"} - feature_names)
    if missing_features:
        return True, "required columns present", False, f"missing features: {missing_features}"

    return True, "required columns present", True, "x_qn24 and x_qn25 present"


def main() -> None:
    parser = argparse.ArgumentParser(description="Week 06 submission readiness checker")
    parser.add_argument("--project-root", type=Path, default=PROJECT_ROOT)
    args = parser.parse_args()

    report_md = args.project_root / "docs/status_reports/report_03/Project_Status_Report_03_Submission.md"
    report_docx = args.project_root / "docs/status_reports/report_03/Project_Status_Report_03_Submission.docx"
    calibration_table = args.project_root / "outputs/tables/week06_calibration_sensitivity_seed2026.csv"
    change_table = args.project_root / "outputs/tables/week06_feature_importance_change_tracking_seed2026.csv"
    audit_path = args.project_root / "docs/status_reports/report_03/week06_submission_readiness_audit.md"

    require_paths([report_md, report_docx, calibration_table, change_table])

    checks: List[Tuple[str, bool, str]] = []

    md_order_ok, md_order_detail = _validate_md_section_order(report_md.read_text(encoding="utf-8"))
    docx_order_ok, docx_order_detail = _validate_docx_section_order(report_docx)
    checks.append(("section order checks", md_order_ok and docx_order_ok, f"markdown={md_order_detail} | docx={docx_order_detail}"))

    constraints_ok, constraints_detail = _validate_text_constraints(report_md, report_docx)
    checks.append(("report constraints checks", constraints_ok, constraints_detail))

    evidence_ok, evidence_detail = _evidence_path_density_checks(report_md, report_docx)
    checks.append(("evidence path density checks", evidence_ok, evidence_detail))

    cal_ok, cal_detail = _calibration_table_check(calibration_table)
    checks.append(("calibration table existence", cal_ok, cal_detail))

    change_ok, change_detail, bullying_ok, bullying_detail = _change_tracking_check(change_table)
    checks.append(("importance change tracking table existence", change_ok, change_detail))
    checks.append(("x_qn24 and x_qn25 presence", bullying_ok, bullying_detail))

    overall_ok = all(flag for _, flag, _ in checks)

    lines: List[str] = []
    lines.append("# Week 06 Submission Readiness Audit")
    lines.append("")
    for name, flag, detail in checks:
        lines.append(f"- {'PASS' if flag else 'FAIL'}: {name}")
        lines.append(f"  - Detail: {detail}")
    lines.append("")
    lines.append(f"Overall: {'PASS' if overall_ok else 'FAIL'}")

    audit_path.parent.mkdir(parents=True, exist_ok=True)
    audit_path.write_text("\n".join(lines), encoding="utf-8")

    if overall_ok:
        print(json.dumps({"status": "ok"}, indent=2))
        return

    print(json.dumps({"status": "fail", "audit": str(audit_path)}, indent=2))
    raise SystemExit(1)


if __name__ == "__main__":
    main()
